"""
Document Processor Module
Handles reading and extracting text from DOCX/DOC files
"""
from pathlib import Path
from typing import List, Dict, Any, Optional
from dataclasses import dataclass
import re
from docx import Document
from docx.table import Table


@dataclass
class ProcessedDocument:
    """Represents a processed legal document"""
    filename: str
    filepath: str
    content: str
    tables: List[str]
    metadata: Dict[str, Any]
    raw_paragraphs: List[str]


class DocumentProcessor:
    """Processor for legal documents (DOCX/DOC)"""
    
    SUPPORTED_EXTENSIONS = {'.docx', '.doc'}
    
    def __init__(self, data_dir: str = "./Data"):
        self.data_dir = Path(data_dir)
    
    def list_documents(self) -> List[Path]:
        """List all supported documents in data directory"""
        documents = []
        for ext in self.SUPPORTED_EXTENSIONS:
            documents.extend(self.data_dir.glob(f"*{ext}"))
        return sorted(documents)
    
    def read_docx(self, filepath: Path) -> ProcessedDocument:
        """Read and extract content from a DOCX file"""
        doc = Document(str(filepath))
        
        # Extract paragraphs
        paragraphs = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                paragraphs.append(text)
        
        # Extract tables
        tables_text = []
        for table in doc.tables:
            table_content = self._extract_table(table)
            if table_content:
                tables_text.append(table_content)
        
        # Combine content
        full_content = "\n\n".join(paragraphs)
        if tables_text:
            full_content += "\n\n" + "\n\n".join(tables_text)
        
        return ProcessedDocument(
            filename=filepath.name,
            filepath=str(filepath),
            content=full_content,
            tables=tables_text,
            metadata={},
            raw_paragraphs=paragraphs
        )
    
    def _extract_table(self, table: Table) -> str:
        """Extract text from a table"""
        rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):
                rows.append(" | ".join(cells))
        return "\n".join(rows) if rows else ""
    
    def process_document(self, filepath: Path) -> Optional[ProcessedDocument]:
        """Process a single document"""
        if not filepath.exists():
            print(f"File not found: {filepath}")
            return None
        
        ext = filepath.suffix.lower()
        if ext == '.docx':
            return self.read_docx(filepath)
        elif ext == '.doc':
            # For .doc files, try reading as docx (some work)
            try:
                return self.read_docx(filepath)
            except Exception as e:
                print(f"Cannot read .doc file {filepath.name}: {e}")
                print("Consider converting to .docx format")
                return None
        else:
            print(f"Unsupported format: {ext}")
            return None
    
    def process_all_documents(self) -> List[ProcessedDocument]:
        """Process all documents in the data directory"""
        documents = []
        for filepath in self.list_documents():
            print(f"Processing: {filepath.name}")
            doc = self.process_document(filepath)
            if doc:
                documents.append(doc)
        return documents


if __name__ == "__main__":
    # Test the processor
    processor = DocumentProcessor("./Data")
    docs = processor.list_documents()
    print(f"Found {len(docs)} documents")
    
    if docs:
        # Test with first document
        doc = processor.process_document(docs[0])
        if doc:
            print(f"\nProcessed: {doc.filename}")
            print(f"Content length: {len(doc.content)} chars")
            print(f"Paragraphs: {len(doc.raw_paragraphs)}")
            print(f"Tables: {len(doc.tables)}")
